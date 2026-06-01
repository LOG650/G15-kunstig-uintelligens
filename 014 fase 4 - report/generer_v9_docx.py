#!/usr/bin/env python3
"""
generer_v9_docx.py
Generates EndeligForskningsrapport_G15_v9.docx from rapport_full.md.
Opens v8.docx as base to inherit all heading/paragraph styles, clears body,
then renders the updated markdown content. Prepends title page + TOC.
"""

import re
import sys
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT_DIR = Path(__file__).parent
ANALYSE_RESULTATER = REPORT_DIR.parent / "006 analyse" / "resultater"
V8_FILE = REPORT_DIR / "EndeligForskningsrapport G15 v8.docx"
MD_FILE = REPORT_DIR / "rapport_full.md"
OUT_FILE = REPORT_DIR / "EndeligForskningsrapport_G15_v10.docx"

TITLE_LINE1 = "Maskinlæringsbasert prognose av norsk lakseksportpris"
TITLE_LINE2 = "4–12 uker frem i tid"
AUTHORS = [
    "Alexander Francke Lindløkken",
    "Joakim Bekkevik Gåseland",
    "Carl-Henrik Solli Lilleng",
    "Morten Røgeberg",
]
INSTITUTION = "Høgskolen i Molde"
STUDY_PROGRAM = ""
COURSE = "LOG650 – Logistikk og kunstig intelligens"
DATE = "Mai 2026"

# Map trailing filename → local path for images
IMAGE_MAP = {
    "rapport_modellsammenligning.png": REPORT_DIR / "rapport_modellsammenligning.png",
    "ml_ensemble_prediksjon.png": REPORT_DIR / "ml_ensemble_prediksjon.png",
    "usikkerhet_kalibrering.png": REPORT_DIR / "rapport_ci_kalibrering.png",
    "sarima_residualer.png": REPORT_DIR / "sarima_residualer.png",
    "ml_avansert_bias_korr.png": REPORT_DIR / "rapport_ensemble_bias.png",
    "ml_avansert_shap_h4.png": REPORT_DIR / "ml_avansert_shap_h4.png",
}


# ---------------------------------------------------------------------------
# DOM helpers
# ---------------------------------------------------------------------------

def insert_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    return para


def add_toc_field(doc):
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "[Høyreklikk og velg «Oppdater felt» for å generere innholdsfortegnelsen]"
    )
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2, placeholder, fldChar3])
    return para


# ---------------------------------------------------------------------------
# Inline markdown → runs
# ---------------------------------------------------------------------------

def apply_inline(para, text):
    """Parse **bold**, *italic*, `code` and add runs to para."""
    # Handle escaped markdown in table cells / headings
    pattern = re.compile(r"(\*\*.*?\*\*|\*[^*]+\*|`[^`]+`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            para.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            para.add_run(part)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def build_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(INSTITUTION)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.size = Pt(14)

    doc.add_paragraph()

    for line in (TITLE_LINE1, TITLE_LINE2):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(20)

    doc.add_paragraph()
    doc.add_paragraph()

    for author in AUTHORS:
        p = doc.add_paragraph(author)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(13)

    doc.add_paragraph()

    for line in filter(None, (STUDY_PROGRAM, COURSE)):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph(DATE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(12)

    insert_page_break(doc)


# ---------------------------------------------------------------------------
# TOC page
# ---------------------------------------------------------------------------

def build_toc_page(doc):
    doc.add_heading("Innholdsfortegnelse", level=1)
    doc.add_paragraph()
    add_toc_field(doc)
    insert_page_break(doc)


# ---------------------------------------------------------------------------
# Table renderer
# ---------------------------------------------------------------------------

def parse_md_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_md_table(doc, table_lines):
    rows = parse_md_table(table_lines)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Normal Table"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # Strip markdown bold markers for plain cell text
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            clean = re.sub(r"\*(.*?)\*", r"\1", clean)
            clean = re.sub(r"`(.*?)`", r"\1", clean)
            cell.text = clean
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Image renderer
# ---------------------------------------------------------------------------

def resolve_image(path_str):
    filename = Path(path_str).name
    if filename in IMAGE_MAP and IMAGE_MAP[filename].exists():
        return IMAGE_MAP[filename]
    for candidate in (REPORT_DIR / filename, ANALYSE_RESULTATER / filename):
        if candidate.exists():
            return candidate
    return None


def add_image_para(doc, alt_text, img_path):
    img_file = resolve_image(img_path)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_file and img_file.exists():
        try:
            para.add_run().add_picture(str(img_file), width=Inches(5.5))
        except Exception as e:
            print(f"  [WARN] {img_file.name}: {e}")
            para.add_run(f"[Figur: {alt_text}]").italic = True
    else:
        para.add_run(f"[Figur: {alt_text}]").italic = True
        print(f"  [WARN] Bildefil ikke funnet: {img_path}")
    caption = doc.add_paragraph(alt_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in caption.runs:
        r.italic = True
        r.font.size = Pt(10)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Blockquote renderer (decision guide)
# ---------------------------------------------------------------------------

def add_blockquote(doc, lines):
    """Render blockquote content. If it contains a table, render it as a proper Word table."""
    table_lines = []
    in_table = False
    
    for line in lines:
        stripped = line.lstrip("> ").strip()
        
        # Table detection
        if stripped.startswith("|"):
            table_lines.append(stripped)
            in_table = True
            continue
        
        # If we were in a table and current line is NOT a table line, flush the table
        if in_table and not stripped.startswith("|"):
            add_md_table(doc, table_lines)
            table_lines = []
            in_table = False
        
        if not stripped:
            continue
            
        # Normal blockquote text (or heading)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        
        if stripped.startswith("#"):
            # Sub-heading inside blockquote
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(12)
        else:
            apply_inline(p, stripped)
            for r in p.runs:
                r.font.size = Pt(11)
                
    # Final flush for table if it ended the blockquote
    if in_table:
        add_md_table(doc, table_lines)


# ---------------------------------------------------------------------------
# Main markdown renderer
# ---------------------------------------------------------------------------

def render_markdown(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_lines = []
    in_blockquote = False
    blockquote_lines = []
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.name = "Courier New"
                r.font.size = Pt(9)
            i += 1
            continue

        # Blockquote accumulation
        if line.startswith(">"):
            blockquote_lines.append(line)
            in_blockquote = True
            i += 1
            continue
        if in_blockquote:
            add_blockquote(doc, blockquote_lines)
            blockquote_lines = []
            in_blockquote = False
            # re-process current line (do NOT increment)

        # Table accumulation
        if line.startswith("|"):
            table_lines.append(line)
            in_table = True
            i += 1
            continue
        if in_table:
            add_md_table(doc, table_lines)
            table_lines = []
            in_table = False
            # re-process current line (do NOT increment)

        # Heading
        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Image
        m = re.match(r"^!\[(.+?)\]\((.+?)\)\s*$", line)
        if m:
            add_image_para(doc, m.group(1), m.group(2))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Cm(1)
            num_run = p.add_run(f"{m.group(1)}. ")
            num_run.bold = False
            apply_inline(p, m.group(2))
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Cm(1)
            bullet_run = p.add_run("• ")
            apply_inline(p, m.group(1))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        apply_inline(p, line.strip())
        i += 1

    # Flush
    if in_table:
        add_md_table(doc, table_lines)
    if in_blockquote:
        add_blockquote(doc, blockquote_lines)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    print(f"Åpner v8 som stilbase: {V8_FILE}")
    doc = Document(str(V8_FILE))

    # Preserve the trailing sectPr (page layout / margins) from v8
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    sectPr_copy = deepcopy(sectPr) if sectPr is not None else None

    # Clear ALL existing body content
    for child in list(body):
        body.remove(child)

    # Re-attach section properties so page layout is intact
    if sectPr_copy is not None:
        body.append(sectPr_copy)

    print("Bygger tittelside og innholdsfortegnelse...")
    build_title_page(doc)
    build_toc_page(doc)

    print(f"Renderer markdown: {MD_FILE}")
    md_text = MD_FILE.read_text(encoding="utf-8")
    render_markdown(doc, md_text)

    print(f"Lagrer: {OUT_FILE}")
    doc.save(str(OUT_FILE))
    print("Ferdig.")


if __name__ == "__main__":
    main()
